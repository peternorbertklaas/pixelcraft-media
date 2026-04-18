from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Users\230200072023\OneDrive - CGM\07_Private\Projects\MediaDesign\website\docs\mockup-prompts.docx"

PROJECTS = [
    {
        "heading": "01 · LuxeCommerce",
        "tags": "Tags: E-Commerce · Webdesign",
        "filename": "Dateiname: 01_luxecommerce-mockup.png",
        "prompt": (
            "A sleek MacBook Pro displayed at a slight angle on a dark matte surface. "
            "The screen shows a premium luxury fashion e-commerce website with a dark editorial design — "
            "full-bleed product photography, thin serif typography, off-white and champagne color palette, "
            "large hero image of a high-end leather bag. Cinematic studio lighting with a subtle warm glow "
            "from the left. Background: near-black (#09090F). No reflections. Photorealistic product mockup style."
        ),
    },
    {
        "heading": "02 · MedTech Solutions",
        "tags": "Tags: SaaS · UX/UI Design",
        "filename": "Dateiname: 02_medtech-mockup.png",
        "prompt": (
            "A modern iMac or wide desktop monitor on a minimal white desk, shot from a slightly elevated angle. "
            "The screen displays a clean SaaS dashboard for medical technology — data visualizations, dark sidebar "
            "navigation, teal and white accent colors, clean sans-serif typography, patient analytics charts. "
            "Soft blue ambient lighting. Background: very dark navy. Crisp shadows, no clutter. "
            "Professional tech aesthetic."
        ),
    },
    {
        "heading": "03 · Gastro Group",
        "tags": "Tags: Restaurants · Branding",
        "filename": "Dateiname: 03_gastrogroup-mockup.png",
        "prompt": (
            "A MacBook and iPhone displayed together on a dark concrete surface. Both screens show the same "
            "restaurant website — warm amber and deep brown color palette, full-screen food photography, "
            "elegant serif headlines, online reservation interface visible on the phone. Moody candlelight-style "
            "side lighting from the right. Background: very dark brown-black. Shallow depth of field, "
            "premium hospitality brand aesthetic."
        ),
    },
    {
        "heading": "04 · FinanceOne",
        "tags": "Tags: Fintech · Landing Page",
        "filename": "Dateiname: 04_financeone-mockup.png",
        "prompt": (
            "A floating iPhone 15 Pro in space grey, centered and slightly tilted, on a pure dark background. "
            "The screen shows a sleek fintech landing page — dark UI with deep blue and electric indigo accents, "
            "financial data graphs, bold sans-serif typography, a prominent CTA button in electric blue. "
            "Subtle glow emanating from the screen illuminating the dark surroundings. Background: near-black "
            "with faint indigo gradient. Minimal and trustworthy fintech aesthetic."
        ),
    },
    {
        "heading": "05 · Vitalis Klinik",
        "tags": "Tags: Healthcare · Web Design",
        "filename": "Dateiname: 05_vitalis-mockup.png",
        "prompt": (
            "A MacBook Pro on a light marble surface, screen angled toward viewer. The website shown is a "
            "premium private medical clinic — clean white and soft sage green palette, calm and trustworthy "
            "design, doctor photography, appointment booking interface prominently visible, elegant serif and "
            "clean sans-serif type combination. Soft natural window light from the left. Background: light warm "
            "grey surface. Premium healthcare brand, not clinical — more like a luxury wellness clinic."
        ),
    },
    {
        "heading": "06 · Bauart Studio",
        "tags": "Tags: Architektur · Branding",
        "filename": "Dateiname: 06_bauart-mockup.png",
        "prompt": (
            "A wide-format monitor mounted on a minimal concrete and steel desk, photographed head-on with "
            "slight top-down angle. The screen displays an architecture studio website — dark editorial design, "
            "full-bleed architectural photography (buildings, interiors), raw concrete textures, bold condensed "
            "serif typography in off-white, very minimal navigation. Background: dark charcoal and concrete. "
            "Hard directional lighting creating strong shadows. Brutalist-meets-premium aesthetic."
        ),
    },
]

FOOTER_NOTES = [
    "Alle Bilder sollten auf dem gleichen dunklen Hintergrundton basieren (#09090F oder ähnlich), damit sie in der Work-Section konsistent wirken",
    "Gerät immer leicht angewinkelt oder in einer natürlichen Haltung — kein reines Front-on",
    "Der Bildschirminhalt sollte erkennbar aber nicht zu detailliert sein",
    "Nach Generierung: Bilder als .png oder .webp in /public/pictures/mockups/ ablegen",
]

ORANGE = RGBColor(255, 77, 0)
GREY   = RGBColor(110, 110, 110)
BLACK  = RGBColor(30, 30, 30)
DARK   = RGBColor(50, 50, 50)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_horizontal_rule(doc):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(6)
    run = title_p.add_run("Mockup Image Prompts — Unsere Arbeiten")
    set_font(run, size=20, bold=True, color=BLACK)

    # ── Intro block ───────────────────────────────────────────────────────────
    intro_lines = [
        ("Ziel: ", "6 hochwertige Website-Mockup-Bilder für die Work-Section."),
        ("Format: ", "16:9 oder 3:2, min. 1600×1000px, dunkler Hintergrund bevorzugt."),
        ("Stil: ", "Premium, editorial, dunkles UI auf Gerät, keine Stock-Foto-Ästhetik."),
        ("Empfohlene Tools: ", "Midjourney, DALL·E 3, Adobe Firefly"),
    ]
    for label, value in intro_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r_label = p.add_run(label)
        set_font(r_label, bold=True, size=11, color=DARK)
        r_val = p.add_run(value)
        set_font(r_val, size=11, color=DARK)

    # Spacer + rule before first project
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_horizontal_rule(doc)

    # ── Projects ──────────────────────────────────────────────────────────────
    for project in PROJECTS:
        # Heading
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(14)
        heading_p.paragraph_format.space_after  = Pt(4)
        run = heading_p.add_run(project["heading"])
        set_font(run, size=14, bold=True, color=ORANGE)

        # Tags line
        tags_p = doc.add_paragraph()
        tags_p.paragraph_format.space_before = Pt(0)
        tags_p.paragraph_format.space_after  = Pt(2)
        run = tags_p.add_run(project["tags"])
        set_font(run, size=11, color=GREY)

        # Filename line
        file_p = doc.add_paragraph()
        file_p.paragraph_format.space_before = Pt(0)
        file_p.paragraph_format.space_after  = Pt(6)
        run = file_p.add_run(project["filename"])
        set_font(run, size=11, color=GREY)

        # Prompt label
        prompt_label_p = doc.add_paragraph()
        prompt_label_p.paragraph_format.space_before = Pt(0)
        prompt_label_p.paragraph_format.space_after  = Pt(2)
        run = prompt_label_p.add_run("Prompt:")
        set_font(run, size=11, bold=True, color=DARK)

        # Prompt text (indented, italic)
        prompt_p = doc.add_paragraph()
        prompt_p.paragraph_format.left_indent   = Cm(0.5)
        prompt_p.paragraph_format.space_before  = Pt(0)
        prompt_p.paragraph_format.space_after   = Pt(10)
        run = prompt_p.add_run(project["prompt"])
        set_font(run, size=11, italic=True, color=RGBColor(60, 60, 60))

        # Divider
        add_horizontal_rule(doc)

    # ── Footer: Allgemeine Hinweise ───────────────────────────────────────────
    footer_heading = doc.add_paragraph()
    footer_heading.paragraph_format.space_before = Pt(14)
    footer_heading.paragraph_format.space_after  = Pt(6)
    run = footer_heading.add_run("Allgemeine Hinweise")
    set_font(run, size=13, bold=True, color=BLACK)

    for note in FOOTER_NOTES:
        note_p = doc.add_paragraph(style="List Bullet")
        note_p.paragraph_format.space_before = Pt(2)
        note_p.paragraph_format.space_after  = Pt(2)
        run = note_p.add_run(note)
        set_font(run, size=11, color=DARK)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
